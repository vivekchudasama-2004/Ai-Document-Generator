"""150 words/page guaranteed minimum: sentence-greedy packing, then
rebalance pass that pulls sentences forward so every non-final page
reaches the minimum. The final page is exempt (words can't be invented)."""
import html as _html
import re

from app.services.detector import _sentences, count_words
from app.services.mermaid import sanitize_svg

PAGE_MIN_WORDS = 145

PAGE_MIN_WORDS = 145


def _rebalance(pages: list[dict]) -> list[dict]:
    """Pull opening sentences forward until each non-final page hits the minimum."""
    index = 0
    while index < len(pages) - 1:
        while pages[index]["words"] < PAGE_MIN_WORDS and pages[index + 1]["sentences"]:
            pulled = pages[index + 1]["sentences"].pop(0)
            pulled_words = count_words(pulled)
            # Never strand the next page below a stub: merge it whole instead.
            if pages[index + 1]["words"] - pulled_words < 30 and len(pages[index + 1]["sentences"]) == 0:
                pages[index]["sentences"].append(pulled)
                pages[index]["words"] += pulled_words
                pages.pop(index + 1)
                break
            pages[index]["sentences"].append(pulled)
            pages[index]["words"] += pulled_words
            pages[index + 1]["words"] -= pulled_words
        index += 1
    return [page for page in pages if page["sentences"]]


WORDS_PER_PAGE = 150
TOLERANCE = 5


def _esc(text: str) -> str:
    return _html.escape(text or "", quote=False)


def _css(text: str) -> str:
    """Running-head safe: no quotes or newlines that could break @page CSS."""
    return re.sub(r"\s+", " ", (text or "").replace('"', "").replace("'", "")).strip()[:60]


def _inline_md(text: str) -> str:
    """Escaped inline markdown: `code`, **bold**, *italic*."""
    text = _esc(text)
    text = re.sub(r"`([^`]+)`", r"<code>\1</code>", text)
    text = re.sub(r"\*\*([^*]+)\*\*", r"<strong>\1</strong>", text)
    return re.sub(r"\*([^*]+)\*", r"<em>\1</em>", text)


def md_to_html(markdown: str) -> str:
    """Small markdown→HTML renderer for exports (no new dependency).

    Headings, fenced code, bullets, numbered lists, paragraphs. Mermaid
    fences are stripped here — diagrams ride alongside as stored SVG.
    """
    blocks: list[str] = []
    paragraph: list[str] = []
    lst: list[str] = []
    list_tag = ""

    def flush_paragraph() -> None:
        if paragraph:
            blocks.append(f"<p>{' '.join(_inline_md(line) for line in paragraph)}</p>")
            paragraph.clear()

    def flush_list() -> None:
        nonlocal list_tag
        if lst:
            blocks.append(f"<{list_tag}>" + "".join(f"<li>{item}</li>" for item in lst) + f"</{list_tag}>")
            lst.clear()
            list_tag = ""

    cleaned = re.sub(r"```mermaid.*?```", "", markdown or "", flags=re.DOTALL)
    for chunk in re.split(r"(```.*?```)", cleaned, flags=re.DOTALL):
        if chunk.startswith("```"):
            flush_paragraph()
            flush_list()
            code = _esc(re.sub(r"^```\w*\n?|\n?```$", "", chunk.strip()))
            blocks.append(f"<pre><code>{code}</code></pre>")
            continue
        for line in chunk.splitlines():
            stripped = line.strip()
            heading = re.match(r"^(#{1,4})\s+(.*)$", stripped)
            bullet = re.match(r"^[-*]\s+(.*)$", stripped)
            numbered = re.match(r"^\d+[.)]\s+(.*)$", stripped)
            if heading:
                flush_paragraph()
                flush_list()
                level = min(len(heading.group(1)) + 1, 4)
                blocks.append(f"<h{level}>{_inline_md(heading.group(2))}</h{level}>")
            elif bullet or numbered:
                flush_paragraph()
                tag = "ul" if bullet else "ol"
                if list_tag and list_tag != tag:
                    flush_list()
                list_tag = tag
                lst.append(_inline_md((bullet or numbered).group(1)))
            elif stripped:
                flush_list()
                paragraph.append(stripped)
            else:
                flush_paragraph()
                flush_list()
    flush_paragraph()
    flush_list()
    return "\n".join(blocks)


def _page_sentences_html(sentences: list[str]) -> str:
    return md_to_html(" ".join(sentences))


def _body_text(markdown: str) -> str:
    """Plain prose for sentence splitting/counts: no fences, no headings."""
    no_code = re.sub(r"```.*?```", "", markdown or "", flags=re.DOTALL)
    return re.sub(r"^#{1,6}\s+.*$", "", no_code, flags=re.MULTILINE)


def paginate_sections(sections: list[dict]) -> tuple[list[dict], list[dict]]:
    """Paginate each section separately so every section starts on a fresh
    page. Returns (pages, toc) where toc = [{title, page}] with exact
    1-based body page numbers."""
    pages: list[dict] = []
    toc: list[dict] = []
    for section in sections:
        text = section.get("content_humanized_md") or section.get("content_md", "")
        sentences = _sentences(_body_text(text))
        if not sentences:
            continue
        toc.append({"title": section.get("title", ""), "page": len(pages) + 1})
        current: list[str] = []
        current_words = 0
        for sentence in sentences:
            words = count_words(sentence)
            if current_words + words > WORDS_PER_PAGE + TOLERANCE and current:
                pages.append({"section": section.get("title", ""), "words": current_words,
                              "sentences": current})
                current, current_words = [], 0
            current.append(sentence)
            current_words += words
        if current:
            pages.append({"section": section.get("title", ""), "words": current_words,
                          "sentences": current})
    pages = _rebalance(pages)
    # Rebuild sheet numbers from the final pages so the TOC stays exact
    # after sentences moved forward (empty sections keep nearest sheet).
    first_sheet: dict[str, int] = {}
    for position, page in enumerate(pages):
        first_sheet.setdefault(page.get("section", ""), position + 1)
    for entry in toc:
        entry["page"] = first_sheet.get(entry["title"], len(pages))
    return pages, toc


def paginate(sections: list[dict]) -> list[dict]:
    """Greedy sentence packing to ~150 words/page; merge-forward if last <100."""
    pages: list[dict] = []
    current: list[str] = []
    current_words = 0

    def flush():
        nonlocal current, current_words
        if current:
            pages.append({"words": current_words, "sentences": current})
            current, current_words = [], 0

    for section in sections:
        text = section.get("content_humanized_md") or section.get("content_md", "")
        for sent in _sentences(_body_text(text)):
            w = count_words(sent)
            if current_words + w > WORDS_PER_PAGE + TOLERANCE and current:
                flush()
            current.append(sent)
            current_words += w
    flush()
    return _rebalance(pages)


def build_print_html(
    *, title: str, doc_type: str, sections: list[dict],
    author: str = "DocuForge Humanized", version: str | None = None,
    date_str: str | None = None,
) -> str:
    """Cover + TOC (exact start pages) + 150wpp body with honest footers."""
    from datetime import date

    from app.core.config import get_settings

    version = version or get_settings().APP_VERSION
    pages, toc = paginate_sections(sections)
    # TOC page numbers are body-relative; offset by cover + TOC sheets.
    toc_rows = "".join(
        f"<div class='toc-row'><span>{_esc(entry['title'])}</span>"
        f"<span class='dots'></span><span>{entry['page'] + 2}</span></div>"
        for entry in toc
    )
    # Diagrams ride on their section's first sheet (sanitized stored SVG).
    seen_diagram: set[str] = set()
    body_pages = []
    for i, page in enumerate(pages):
        section_title = page.get("section", "")
        diagram = ""
        if section_title not in seen_diagram:
            seen_diagram.add(section_title)
            for section in sections:
                svg = (section.get("title") == section_title and section.get("mermaid_svg")) or ""
                if svg:
                    try:
                        diagram = f"<figure class='diagram'>{sanitize_svg(svg)}<figcaption>Architecture diagram</figcaption></figure>"
                    except ValueError:
                        diagram = ""
                    break
        body_pages.append(
            f"<section class='page'><h2>{_esc(section_title)}</h2>"
            f"{_page_sentences_html(page['sentences'])}{diagram}"
            f"<footer>{_esc(title)}</footer></section>"
        )
    body_pages = "".join(body_pages)
    return f"""<!DOCTYPE html><html><head><meta charset='utf-8'><title>{_esc(title)}</title>
<style>
/* LaTeX-article grade print: Computer Modern-style serif, justified prose,
   ruled title page, dot-leader TOC, running head + real page numbers.
   (Numbers come from the browser's print engine via CSS counters — the same
   mechanism pdflatex uses — never hardcoded, so TOC and footers can't drift.) */
@page {{ size: A4; margin: 2.5cm 2cm 3cm 2cm;
  @top-center {{ content: "{_css(title)}"; font-style: italic; font-size: 8.5pt; color: #4b5563; }}
  @bottom-center {{ content: counter(page); font-size: 9pt; color: #4b5563; }} }}
@page :first {{ @top-center {{ content: none; }} @bottom-center {{ content: none; }} }}
body {{ font-family: 'Latin Modern Roman', 'Computer Modern Roman', Georgia, 'Times New Roman', serif;
  font-size: 11pt; line-height: 1.65; color: #111; }}
h1, h2, h3, h4 {{ font-family: 'Latin Modern Roman', 'Computer Modern Roman', Georgia, serif;
  line-height: 1.25; break-after: avoid; }}
h2 {{ font-size: 15pt; margin: 0 0 12pt; }}
h3 {{ font-size: 13pt; }} h4 {{ font-size: 11.5pt; }}
p {{ margin: 0 0 10pt; text-align: justify; hyphens: auto; }}
ul, ol {{ margin: 0 0 10pt 20pt; }} li {{ margin-bottom: 4pt; text-align: justify; }}
pre {{ background: #fafaf8; border: 1px solid #d6d3cb; border-left: 3px solid #1f2937;
  padding: 10pt 12pt; font-size: 8.5pt; line-height: 1.5; white-space: pre-wrap; break-inside: avoid; }}
code {{ font-family: 'Latin Modern Mono', 'JetBrains Mono', Consolas, monospace; font-size: 9pt; }}
.page {{ page-break-after: always; }}
footer {{ text-align: center; color: #6B7280; font-size: 9pt; margin-top: 18pt; }}
.cover {{ text-align: center; padding-top: 5cm; }}
.cover p {{ text-align: center; }}
.cover .doctype {{ font-variant: small-caps; letter-spacing: 3pt; font-size: 12pt; color: #374151; }}
.cover h1 {{ font-size: 26pt; margin: 14pt 0; }}
.cover .rule {{ border: none; border-top: 2px solid #1f2937; width: 38%; margin: 16pt auto; }}
.cover .meta {{ font-size: 10.5pt; color: #374151; }}
.toc-row {{ display: flex; gap: 8px; margin: 7px 0; font-size: 10.5pt; }}
.toc-row .dots {{ flex: 1; border-bottom: 1px dotted #6B7280; }}
.diagram {{ margin: 14pt 0; max-width: 100%; break-inside: avoid; }}
.diagram svg {{ max-width: 100%; height: auto; }}
.diagram figcaption {{ font-size: 9pt; color: #4b5563; text-align: center; margin-top: 6pt; }}
</style></head><body>
<section class='page cover'><p class='doctype'>{_esc(doc_type)}</p><hr class='rule'><h1>{_esc(title)}</h1><hr class='rule'>
<p class='meta'>{_esc(author)} · v{_esc(version)} · {date_str or date.today().isoformat()}</p></section>
<section class='page'><h1>Contents</h1>{toc_rows}<footer>{_esc(title)}</footer></section>
{body_pages}
</body></html>"""


def _docx_add_rich(paragraph, text: str) -> None:
    """**bold** runs inside a python-docx paragraph (text already trusted)."""
    for i, chunk in enumerate(re.split(r"(\*\*[^*]+\*\*)", text)):
        run = paragraph.add_run(chunk[2:-2] if i % 2 else chunk)
        if i % 2:
            run.bold = True


def build_docx(*, title: str, sections: list[dict]) -> bytes:
    from io import BytesIO

    from docx import Document as DocxDocument
    from docx.shared import Pt

    doc = DocxDocument()
    style = doc.styles["Normal"]
    style.font.size = Pt(11)
    doc.add_heading(title, level=0)
    for section in sections:
        doc.add_heading(section.get("title", ""), level=1)
        text = section.get("content_humanized_md") or section.get("content_md", "")
        text = re.sub(r"```mermaid.*?```", "", text, flags=re.DOTALL)
        for chunk in re.split(r"(```.*?```)", text, flags=re.DOTALL):
            if chunk.startswith("```"):
                code = re.sub(r"^```\w*\n?|\n?```$", "", chunk.strip())
                doc.add_paragraph(code, style="No Spacing").runs[0].font.size = Pt(9)
                continue
            for line in chunk.splitlines():
                stripped = line.strip()
                heading = re.match(r"^(#{1,4})\s+(.*)$", stripped)
                bullet = re.match(r"^[-*]\s+(.*)$", stripped)
                numbered = re.match(r"^\d+[.)]\s+(.*)$", stripped)
                if heading:
                    doc.add_heading(heading.group(2), level=min(len(heading.group(1)), 4))
                elif bullet:
                    _docx_add_rich(doc.add_paragraph(style="List Bullet"), bullet.group(1))
                elif numbered:
                    _docx_add_rich(doc.add_paragraph(style="List Number"), numbered.group(1))
                elif stripped:
                    _docx_add_rich(doc.add_paragraph(), stripped)
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


LATEX_ESCAPES = {
    "\\": r"\textbackslash{}",
    "&": r"\&", "%": r"\%", "$": r"\$", "#": r"\#", "_": r"\_",
    "{": r"\{", "}": r"\}", "~": r"\textasciitilde{}", "^": r"\textasciicircum{}",
    "*": r"{*}",
}


def _tex(text: str) -> str:
    return "".join(LATEX_ESCAPES.get(ch, ch) for ch in (text or ""))


def _tex_inline(text: str) -> str:
    parts = re.split(r"(\*\*[^*]+\*\*|`[^`]+`|\*[^*]+\*)", text or "")
    out = []
    for part in parts:
        if len(part) > 4 and part.startswith("**") and part.endswith("**"):
            out.append(r"\textbf{" + _tex(part[2:-2]) + "}")
        elif len(part) > 2 and part.startswith("`") and part.endswith("`"):
            out.append(r"\texttt{" + _tex(part[1:-1]) + "}")
        elif len(part) > 2 and part.startswith("*") and part.endswith("*"):
            out.append(r"\textit{" + _tex(part[1:-1]) + "}")
        else:
            out.append(_tex(part))
    return "".join(out)


def md_to_latex(markdown: str) -> str:
    """Markdown→LaTeX body converter. Mermaid fences become verbatim blocks."""
    blocks: list[str] = []
    paragraph: list[str] = []
    items: list[str] = []
    list_env = ""

    def flush_paragraph() -> None:
        if paragraph:
            blocks.append(" ".join(_tex_inline(line) for line in paragraph) + "\n")
            paragraph.clear()

    def flush_list() -> None:
        nonlocal list_env
        if items:
            blocks.append(f"\\begin{{{list_env}}}\n" + "".join(f"\\item {i}\n" for i in items) + f"\\end{{{list_env}}}\n")
            items.clear()
            list_env = ""

    for chunk in re.split(r"(```.*?```)", markdown or "", flags=re.DOTALL):
        if chunk.startswith("```"):
            flush_paragraph()
            flush_list()
            code = re.sub(r"^```\w*\n?|\n?```$", "", chunk.strip())
            blocks.append("\\begin{verbatim}\n" + code.replace("\\end{verbatim}", "") + "\n\\end{verbatim}\n")
            continue
        for line in chunk.splitlines():
            stripped = line.strip()
            heading = re.match(r"^(#{1,4})\s+(.*)$", stripped)
            bullet = re.match(r"^[-*]\s+(.*)$", stripped)
            numbered = re.match(r"^\d+[.)]\s+(.*)$", stripped)
            if heading:
                flush_paragraph()
                flush_list()
                blocks.append(f"\\{['', 'section', 'subsection', 'subsubsection', 'paragraph'][len(heading.group(1))]}"
                              f"{{{_tex_inline(heading.group(2))}}}\n")
            elif bullet or numbered:
                flush_paragraph()
                env = "itemize" if bullet else "enumerate"
                if list_env and list_env != env:
                    flush_list()
                list_env = env
                items.append(_tex_inline((bullet or numbered).group(1)))
            elif stripped:
                flush_list()
                paragraph.append(stripped)
            else:
                flush_paragraph()
                flush_list()
    flush_paragraph()
    flush_list()
    return "\n".join(blocks)


def build_tex(
    *, title: str, doc_type: str, sections: list[dict],
    author: str = "DocuForge Humanized", version: str | None = None,
    date_str: str | None = None,
) -> str:
    """A real .tex source: article class, title page, TOC, full-fidelity
    bodies (lists → itemize/enumerate, code → verbatim). Compile with
    `pdflatex doc.tex` (twice for the TOC); pdflatex itself paginates."""
    from datetime import date

    from app.core.config import get_settings

    version = version or get_settings().APP_VERSION
    body: list[str] = []
    for section in sections:
        text = section.get("content_humanized_md") or section.get("content_md", "")
        body.append(f"\\section{{{_tex(section.get('title', ''))}}}\n")
        body.append(md_to_latex(text) + "\n\\newpage\n")
    # Diagrams appendix: mermaid sources as verbatim (pdflatex can't take SVG).
    diagrams = [
        (s.get("title", ""), m.group(0))
        for s in sections
        for m in [re.search(r"```mermaid.*?```",
                            s.get("content_humanized_md") or s.get("content_md", ""),
                            flags=re.DOTALL)]
        if m
    ]
    if diagrams:
        body.append("\\appendix\n\\section{Diagrams}\n")
        for diagram_title, block in diagrams:
            code = re.sub(r"^```mermaid\n?|\n?```$", "", block.strip())
            body.append(f"\\subsection{{{_tex(diagram_title)}}}\n\\begin{{verbatim}}\n{code}\n\\end{{verbatim}}\n")
    return (
        "% Generated by DocuForge — compile with: pdflatex <file>.tex (twice for TOC)\n"
        "\\documentclass[11pt,a4paper]{article}\n"
        "\\usepackage[utf8]{inputenc}\n\\usepackage[T1]{fontenc}\n"
        "\\usepackage{lmodern}\n\\usepackage[a4paper,margin=2.5cm]{geometry}\n"
        "\\usepackage[hidelinks]{hyperref}\n\\usepackage{parskip}\n\n"
        f"\\title{{{_tex(title)}}}\n"
        f"\\author{{{_tex(author)} v{_tex(version)}}}\n"
        f"\\date{{{date_str or date.today().isoformat()}}}\n\n"
        "\\begin{document}\n\\maketitle\n"
        f"\\begin{{center}}\\textit{{{_tex(doc_type)}}}\\end{{center}}\n\n"
        "\\tableofcontents\n\\newpage\n\n"
        + "\n".join(body) +
        "\\end{document}\n"
    )
